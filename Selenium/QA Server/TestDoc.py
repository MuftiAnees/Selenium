import os
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from time import sleep
from colorama import Fore, Back, Style
from CommonScripts.login import *
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.shared import Inches

PATH = "D:\Coding\Python\Selenium\chromedriver.exe"
document = Document()


driver.maximize_window()
login()
os.system('cls')
print(Fore.GREEN+'------------------------------------------------------------Test Started!------------------------------------------------------------')
print(Fore.RESET)

links = ["http://172.16.10.4:8082/jw/web/userview/overtimeAllowance/overtimeallowanceUV/_/standardHours",
         "http://172.16.10.4:8082/jw/web/userview/overtimeAllowance/overtimeallowanceUV/_/rateConfig",
         "http://172.16.10.4:8082/jw/web/userview/overtimeAllowance/overtimeallowanceUV/_/Monthly"]
total_number_of_links = len(links)
linkNumber = 0


document.add_heading('Module Report', 0)

while linkNumber != total_number_of_links:
    driver.get(links[linkNumber])
    document.add_heading(driver.title, level=1)
    print(Fore.YELLOW+driver.title)
    print(Fore.RESET)
    sleep(2)

    # ----------------------------------------------------------SearchBar----------------------------------------------------------
    MandatorySearchBar = 0
    try:
        searchbar = driver.find_element(By.ID, 'st')
        print(Fore.GREEN+'Searchbar Present')
        print(Fore.RESET)
        p = document.add_paragraph('Searchbar Present')
        MandatorySearchBar += 1
    except NoSuchElementException:
        print(Fore.RED+'Searchbar Missing')
        p = document.add_paragraph('Searchbar Missing')
        print(Fore.RESET)

    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="Reset"]')
        p = document.add_paragraph('Searchbar Reset Button Present')
        print(Fore.GREEN+'Searchbar Reset Button Present')
        print(Fore.RESET)
        MandatorySearchBar += 1
    except NoSuchElementException:
        p = document.add_paragraph('Searchbar Reset Button Missing')
        print(Fore.RED+'Searchbar Reset Button Missing')
        print(Fore.RESET)

    # ----------------------------------------------------------SearchBar----------------------------------------------------------

    # ----------------------------------------------------------Group Bar----------------------------------------------------------
    MandatoryGroup = 0
    try:
        searchbar = driver.find_element(By.CLASS_NAME, 'k-grouping-header')
        p = document.add_paragraph('Group bar Present')
        print(Fore.GREEN+'Group bar Present')
        print(Fore.RESET)
        MandatoryGroup = 1
    except NoSuchElementException:
        p = document.add_paragraph('Group bar Missing')
        print(Fore.RED+'Group bar Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Group Bar----------------------------------------------------------

    # ----------------------------------------------------------C.R.U.D----------------------------------------------------------
    print(Fore.BLUE+'C.R.U.D Icons Test started')
    document.add_heading('C.R.U.D Icons Test', level=2)
    MandatoryCRUD = 0

    # ----------------------------------------------------------Reset State----------------------------------------------------------

    try:
        searchbar = driver.find_element(By.ID, 'reloadGridState')
        p = document.add_paragraph('Reset State Button Present')
        print(Fore.GREEN+'Reset State Button Present')
        print(Fore.RESET)
        MandatoryCRUD += 1

    except NoSuchElementException:
        p = document.add_paragraph('Reset State Button Missing')
        print(Fore.RED+'Reset State Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Reset State----------------------------------------------------------
    # ----------------------------------------------------------CSV----------------------------------------------------------
    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="CSV"]')
        p = document.add_paragraph('CSV Button Present')
        print(Fore.GREEN+'CSV Button Present')
        print(Fore.RESET)
    except NoSuchElementException:
        p = document.add_paragraph('CSV Button Missing')
        print(Fore.RED+'CSV Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------CSV----------------------------------------------------------
    # ----------------------------------------------------------Save State----------------------------------------------------------
    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="Save State"]')
        p = document.add_paragraph('Save State Button Present')
        print(Fore.GREEN+'Save State Button Present')
        print(Fore.RESET)
        MandatoryCRUD += 1

    except NoSuchElementException:
        p = document.add_paragraph('Save State Button Missing')
        print(Fore.RED+'Save State Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Save State----------------------------------------------------------

    # ----------------------------------------------------------Load State----------------------------------------------------------
    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="Load State"]')
        p = document.add_paragraph('Load State Button Present')
        print(Fore.GREEN+'Load State Button Present')
        print(Fore.RESET)
        MandatoryCRUD += 1

    except NoSuchElementException:
        p = document.add_paragraph('Load State Button Missing')
        print(Fore.RED+'Load State Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Load State----------------------------------------------------------
    # ----------------------------------------------------------Add----------------------------------------------------------
    add_count = 0
    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="Add"]')
        add_count += 1
    except NoSuchElementException:
        add_count += 0

    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="add"]')
        add_count += 1
    except NoSuchElementException:
        add_count += 0
    # ----------------------------------------------------------Add----------------------------------------------------------
    if add_count == 1:
        print(Fore.GREEN+'Add Button Present')
        p = document.add_paragraph('Add Button Present')
        print(Fore.RESET)
    else:
        print(Fore.RED+'Add Button Missing')
        p = document.add_paragraph('Add Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Dormant----------------------------------------------------------
    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="Dormant"]')
        p = document.add_paragraph('Dormant Button Present')
        print(Fore.GREEN+'Dormant Button Present')
        print(Fore.RESET)
        MandatoryCRUD += 1

    except NoSuchElementException:
        p = document.add_paragraph('Dormant Button Missing')
        print(Fore.RED+'Dormant Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Dormant----------------------------------------------------------

    # ----------------------------------------------------------Delete----------------------------------------------------------
    try:
        searchbar = driver.find_element(By.XPATH, '//*[@title="Delete"]')
        p = document.add_paragraph('Delete Button Present')
        print(Fore.GREEN+'Delete Button Present')
        print(Fore.RESET)
    except NoSuchElementException:
        p = document.add_paragraph('Delete Button Missing')
        print(Fore.RED+'Delete Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Delete----------------------------------------------------------

    # ----------------------------------------------------------Edit----------------------------------------------------------

    try:
        formEdit = driver.find_element(
            By.CLASS_NAME, "k-button k-button-icontext Edit k-grid-Edit")
        p = document.add_paragraph('Form Edit Button Present')
        print(Fore.GREEN+'Form Edit Button Present')
        print(Fore.RESET)
    except NoSuchElementException:
        p = document.add_paragraph('Form Edit Button Missing')
        print(Fore.YELLOW+'Form Edit Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Edit----------------------------------------------------------

    # ----------------------------------------------------------Status Update----------------------------------------------------------

    try:
        formEdit = driver.find_element(
            By.CLASS_NAME, "k-button k-button-icontext Status-Update k-grid-StatusUpdate")
        p = document.add_paragraph('Status Update Button Present')
        print(Fore.GREEN+'Status Update Button Present')
        print(Fore.RESET)
    except NoSuchElementException:
        p = document.add_paragraph('Status Update Button Missing')
        print(Fore.YELLOW+'Status Update Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------Status update----------------------------------------------------------

    # ----------------------------------------------------------View----------------------------------------------------------

    try:
        formEdit = driver.find_element(
            By.CLASS_NAME, "k-button k-button-icontext View k-grid-View")
        p = document.add_paragraph('View Button Present')
        print(Fore.GREEN+'View Button Present')
        print(Fore.RESET)
    except NoSuchElementException:
        p = document.add_paragraph('View Button Missing')
        print(Fore.YELLOW+'View Button Missing')
        print(Fore.RESET)
    # ----------------------------------------------------------View----------------------------------------------------------
    # ----------------------------------------------------------Grid----------------------------------------------------------
    sleep(5)
    try:
        searchbar = driver.find_element(By.CSS_SELECTOR, '[role="gridcell"]')
        p = document.add_paragraph('Grid data present')
        print(Fore.GREEN+'Grid data present')
        print(Fore.RESET)
        MandatoryGroup = 1
    except NoSuchElementException:
        p = document.add_paragraph('Grid is empty!')
        print(Fore.RED+'Grid is empty!')
        print(Fore.RESET)
# ----------------------------------------------------------Grid----------------------------------------------------------

    print(Fore.CYAN+'Final Result:')
    document.add_heading('Final Result', level=1)
    print(Fore.RESET)

    if MandatoryCRUD < 4:
        p = document.add_paragraph('C.R.U.D Test Failed')
        print(Fore.RED+'C.R.U.D Test Failed')
        print(Fore.RESET)
    else:
        p = document.add_paragraph('C.R.U.D Test Passed')
        print(Fore.GREEN+'C.R.U.D Test Passed')
        print(Fore.RESET)

    if MandatorySearchBar == 2:
        p = document.add_paragraph('Searchbar Test Passed')
        print(Fore.GREEN+'Searchbar Test Passed')
        print(Fore.RESET)
    else:
        p = document.add_paragraph('Searchbar Test Failed')
        print(Fore.RED+'Searchbar Test Failed')
        print(Fore.RESET)

    if MandatoryGroup == 1:
        p = document.add_paragraph('Group Test Passed')
        print(Fore.GREEN+'Group Test Passed')
        print(Fore.RESET)
    else:
        p = document.add_paragraph('Group Test Failed')
        print(Fore.RED+'Group Test Failed')
        print(Fore.RESET)

    linkNumber += 1
print(Fore.CYAN+'Test Completed')
print(Fore.RESET)
document.save('Result.docx')
sleep(2)

os.system('start Result.docx')
sleep(3600)
